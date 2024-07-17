from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Add Title
doc.add_heading('Pomodoro App Requirements Document', 0)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph("The purpose of this document is to outline the requirements, design, and implementation plan for a Pomodoro productivity tracking application. The application will help users track their work sessions and breaks, providing statistics to improve productivity.")
doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph("This application will consist of a command-line interface (CLI) for initiating Pomodoro sessions, a web interface for monitoring and statistics, and deployment in a Docker container for ease of deployment and scaling.")

# User Personas
doc.add_heading('2. User Personas', level=1)
doc.add_heading('2.1 John Doe - The Busy Professional', level=2)
doc.add_paragraph("Age: 35\nOccupation: Marketing Manager\nNeeds: Track productivity, improve work habits, visualize time spent on tasks\nTech Skill: Intermediate")
doc.add_heading('2.2 Jane Smith - The Student', level=2)
doc.add_paragraph("Age: 22\nOccupation: University Student\nNeeds: Manage study sessions, balance study and break time, analyze study patterns\nTech Skill: Beginner")

# User Stories
doc.add_heading('3. User Stories', level=1)
doc.add_paragraph("1. As a user, I want to start a Pomodoro session from the CLI, so that I can begin working on a task.\n2. As a user, I want to be notified when a Pomodoro session ends, so that I can take a break.\n3. As a user, I want to view my productivity statistics on a web page, so that I can analyze my work patterns.\n4. As a user, I want to track my completed Pomodoro sessions, so that I can see my progress over time.")

# Functional Requirements
doc.add_heading('4. Functional Requirements', level=1)
doc.add_paragraph("1. Start a Pomodoro Session (CLI)\n    - Input: Task description, session duration\n    - Output: Timer starts, notification when session ends\n2. End a Pomodoro Session (CLI)\n    - Input: Session end command\n    - Output: Session stops, log entry created\n3. View Statistics (Web)\n    - Input: User login\n    - Output: Display charts and graphs of completed sessions, break times, and productivity trends\n4. User Authentication (Web)\n    - Input: Username, password\n    - Output: User login, session management\n5. Docker Deployment\n    - Input: Dockerfile and Docker Compose configuration\n    - Output: Deployable Docker containers for CLI and web interface")

# Non-Functional Requirements
doc.add_heading('5. Non-Functional Requirements', level=1)
doc.add_paragraph("1. Usability\n    - Simple and intuitive CLI commands\n    - User-friendly web interface\n2. Performance\n    - Minimal latency in command execution\n    - Quick loading of web pages\n3. Reliability\n    - Robust error handling in CLI and web interface\n    - Consistent uptime for Docker deployment\n4. Scalability\n    - Ability to handle multiple users concurrently\n    - Efficient resource usage in Docker containers")

# System Architecture
doc.add_heading('6. System Architecture', level=1)
doc.add_heading('6.1 System Components', level=2)
doc.add_paragraph("CLI Application: Python-based command-line interface for managing Pomodoro sessions.\nWeb Application: Flask-based web interface for displaying statistics.\nDatabase: SQLite for storing user data and session logs.\nDocker: Containerization of both CLI and web application for deployment.")

# UML Diagrams
doc.add_heading('6.2 UML Diagrams', level=2)
doc.add_heading('Use Case Diagram', level=3)
doc.add_picture("use_case_diagram.png", width=Inches(5))
doc.add_heading('Class Diagram', level=3)
doc.add_picture("class_diagram.png", width=Inches(5))
doc.add_heading('Sequence Diagram', level=3)
doc.add_picture("sequence_diagram.png", width=Inches(5))
doc.add_heading('Deployment Diagram', level=3)
doc.add_picture("deployment_diagram.png", width=Inches(5))
doc.add_heading('System Diagram', level=3)
doc.add_picture("system_diagram.png", width=Inches(5))

# Deployment Diagram
doc.add_heading('7. Deployment Diagram', level=1)
doc.add_paragraph("Deployment diagram showing the system architecture and components.")

# System Diagram
doc.add_heading('8. System Diagram', level=1)
doc.add_paragraph("System diagram illustrating the relationship between the CLI, web application, and database.")

# To-Do List and Priority Order
doc.add_heading('9. To-Do List and Priority Order', level=1)
doc.add_paragraph("1. Define User Stories and Personas\n    - Priority: High\n    - Status: Completed\n2. Gather Functional and Non-Functional Requirements\n    - Priority: High\n    - Status: Completed\n3. Create UML Diagrams\n    - Priority: High\n    - Status: In Progress\n4. Develop CLI Application\n    - Priority: High\n    - Status: Pending\n5. Develop Web Application\n    - Priority: Medium\n    - Status: Pending\n6. Set Up Database\n    - Priority: Medium\n    - Status: Pending\n7. Implement User Authentication\n    - Priority: Medium\n    - Status: Pending\n8. Containerize Applications with Docker\n    - Priority: Medium\n    - Status: Pending\n9. Write Tests for CLI and Web Application\n    - Priority: Medium\n    - Status: Pending\n10. Deploy Applications\n    - Priority: Medium\n    - Status: Pending\n11. Debug and Test Applications\n    - Priority: High\n    - Status: Pending\n12. Launch and Monitor Application\n    - Priority: High\n    - Status: Pending")

# Save Document
doc.save("Pomodoro_App_Requirements_Document1.docx")
